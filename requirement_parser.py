"""
requirement_parser.py
---------------------
Handles parsing and extraction of text from various document formats.
Supports: PDF, DOCX, TXT, and plain text input.
"""

import io
import re


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        return clean_text("\n\n".join(text_parts))
    except ImportError:
        return "Error: pypdf library not installed. Run: pip install pypdf"
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name.split()[-1]
                    prefix = "#" * int(level) if level.isdigit() else "#"
                    text_parts.append(f"{prefix} {para.text}")
                else:
                    text_parts.append(para.text)

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_rows.append(row_text)
            if table_rows:
                text_parts.append("\nTable:\n" + "\n".join(table_rows) + "\n")

        return clean_text("\n\n".join(text_parts))
    except ImportError:
        return "Error: python-docx library not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"


def parse_txt(file_bytes: bytes) -> str:
    """Parse a plain text file, trying UTF-8 then falling back to latin-1."""
    try:
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        return clean_text(text)
    except Exception as e:
        return f"Error parsing TXT: {str(e)}"


def parse_uploaded_file(uploaded_file) -> tuple[str, str]:
    """Parse a Streamlit UploadedFile based on its extension."""
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith('.pdf'):
        return parse_pdf(file_bytes), 'PDF'
    elif filename.endswith('.docx'):
        return parse_docx(file_bytes), 'DOCX'
    elif filename.endswith('.txt'):
        return parse_txt(file_bytes), 'TXT'
    elif filename.endswith('.md'):
        return parse_txt(file_bytes), 'Markdown'
    else:
        try:
            text = file_bytes.decode('utf-8')
            return clean_text(text), 'Text'
        except Exception:
            return "Error: Unsupported file format.", 'Unknown'


def clean_text(text: str) -> str:
    """Clean and normalize extracted text, collapsing excess whitespace."""
    if not text:
        return ""

    lines = text.split('\n')
    cleaned_lines = []
    empty_count = 0

    for line in lines:
        stripped = line.rstrip()
        if stripped == '':
            empty_count += 1
            if empty_count <= 2:
                cleaned_lines.append('')
        else:
            empty_count = 0
            cleaned_lines.append(stripped)

    cleaned = '\n'.join(cleaned_lines).strip()
    # Remove non-printable characters except common whitespace
    cleaned = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\x80-\xFF]', '', cleaned)
    return cleaned


def truncate_for_api(text: str, max_tokens: int = 12000) -> str:
    """
    Truncate text to fit within Groq context limits.
    Groq's free-tier models have smaller effective context than larger hosted
    models, so we keep a conservative budget (~4 chars per token).
    """
    max_chars = max_tokens * 4

    if len(text) <= max_chars:
        return text

    truncated = text[:max_chars]
    last_period = truncated.rfind('.')
    if last_period > max_chars * 0.9:
        truncated = truncated[:last_period + 1]

    return truncated + "\n\n[Document truncated for processing. Showing first ~12,000 tokens]"


def estimate_token_count(text: str) -> int:
    """Rough estimation of token count (4 chars per token)."""
    return len(text) // 4


def detect_document_type(text: str) -> str:
    """Auto-detect the type of requirements document from its content."""
    text_lower = text.lower()

    if any(m in text_lower for m in ['software requirements specification', 'srs', 'system requirements']):
        return "Software Requirements Specification (SRS)"
    elif any(m in text_lower for m in ['business requirements document', 'brd', 'business requirement']):
        return "Business Requirements Document (BRD)"
    elif any(m in text_lower for m in ['as a ', 'as an ', 'user story', 'acceptance criteria']):
        return "User Stories"
    elif any(m in text_lower for m in ['product requirements', 'prd', 'product requirement']):
        return "Product Requirements Document (PRD)"
    elif any(m in text_lower for m in ['functional specification', 'functional requirement']):
        return "Functional Specification Document"
    else:
        return "Requirements Document"
