"""
utils.py
--------
Utility functions for export, formatting, and dashboard metric helpers.
Handles PDF / DOCX / Markdown / JSON export of architecture reports.
"""

import json
import io
from datetime import datetime


def generate_markdown_report(architecture: dict) -> str:
    """Generate a comprehensive Markdown report from architecture data."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    md = f"""# Software Architecture Design Report
*Generated on: {now}*

---

## Project Summary

{architecture.get('project_summary', 'No summary available')}

---

## Functional Requirements

"""
    fr_list = architecture.get('functional_requirements', [])
    md += "".join(f"- {r}\n" for r in fr_list) if fr_list else "*None identified*\n"

    md += "\n---\n\n## Non-Functional Requirements\n\n"
    nfr_list = architecture.get('non_functional_requirements', [])
    md += "".join(f"- {r}\n" for r in nfr_list) if nfr_list else "*None identified*\n"

    md += f"""
---

## Recommended Architecture

{architecture.get('recommended_architecture', 'N/A')}

**Architecture Type:** `{architecture.get('architecture_type', 'N/A')}`

---

## System Components

"""
    for comp in architecture.get('components', []):
        if isinstance(comp, dict):
            md += f"### {comp.get('name', 'Unknown Component')}\n\n"
            md += f"- **Type:** {comp.get('type', 'N/A')}\n"
            md += f"- **Technology:** {comp.get('technology', 'N/A')}\n"
            md += f"- **Description:** {comp.get('description', 'N/A')}\n"
            resp = comp.get('responsibilities', [])
            if resp:
                md += "\n**Responsibilities:**\n" + "".join(f"  - {r}\n" for r in resp)
            deps = comp.get('dependencies', [])
            if deps:
                md += f"\n**Dependencies:** {', '.join(deps)}\n"
            md += "\n"

    md += "---\n\n## Database Design\n\n"
    db_design = architecture.get('database_design', {})
    if isinstance(db_design, dict):
        primary_db = db_design.get('primary_database', {})
        if isinstance(primary_db, dict):
            md += f"### Primary Database: {primary_db.get('technology', 'N/A')}\n\n"
            md += f"- **Type:** {primary_db.get('type', 'N/A')}\n"
            md += f"- **Justification:** {primary_db.get('justification', 'N/A')}\n\n"
            entities = primary_db.get('entities', [])
            for entity in entities:
                if isinstance(entity, dict):
                    md += f"#### {entity.get('name', 'Unknown')}\n"
                    fields = entity.get('fields', [])
                    if fields:
                        md += "| Field | Type |\n|-------|------|\n"
                        for field in fields:
                            parts = field.split(':') if ':' in field else [field, 'unknown']
                            md += f"| {parts[0]} | {parts[1] if len(parts) > 1 else 'N/A'} |\n"
                    md += "\n"
        for db in db_design.get('secondary_databases', []):
            if isinstance(db, dict):
                md += f"- **{db.get('purpose', 'N/A')}:** {db.get('technology', 'N/A')} - {db.get('justification', 'N/A')}\n"
    elif isinstance(db_design, str):
        md += db_design + "\n"

    md += "\n---\n\n## API Design\n\n"
    api_design = architecture.get('api_design', {})
    if isinstance(api_design, dict):
        md += f"- **Style:** {api_design.get('style', 'N/A')}\n"
        md += f"- **Authentication:** {api_design.get('authentication', 'N/A')}\n"
        md += f"- **Versioning:** {api_design.get('versioning_strategy', 'N/A')}\n\n"
        endpoints = api_design.get('endpoints', [])
        if endpoints:
            md += "| Endpoint | Methods | Auth Required | Description |\n"
            md += "|----------|---------|---------------|-------------|\n"
            for ep in endpoints:
                if isinstance(ep, dict):
                    methods = ', '.join(ep.get('methods', []))
                    auth = 'Yes' if ep.get('auth_required') else 'No'
                    md += f"| `{ep.get('resource', 'N/A')}` | {methods} | {auth} | {ep.get('description', 'N/A')} |\n"
    elif isinstance(api_design, str):
        md += api_design + "\n"

    md += "\n---\n\n## Technology Stack Recommendations\n\n"
    tech_stack = architecture.get('tech_stack_recommendations', [])
    if tech_stack:
        md += "| Layer | Technology | Version | Alternatives |\n|-------|-----------|---------|-------------|\n"
        for tech in tech_stack:
            if isinstance(tech, dict):
                alts = ', '.join(tech.get('alternatives', []))
                md += f"| {tech.get('layer', 'N/A')} | **{tech.get('technology', 'N/A')}** | {tech.get('version', 'N/A')} | {alts} |\n"

    md += "\n---\n\n## Design Conflicts & Risks\n\n"
    conflicts = architecture.get('design_conflicts', [])
    if conflicts:
        for c in conflicts:
            if isinstance(c, dict):
                md += f"### [{c.get('severity', 'medium').upper()}] {c.get('type', 'Issue').replace('_', ' ').title()}\n\n"
                md += f"{c.get('description', 'N/A')}\n\n"
                md += f"**Recommendation:** {c.get('recommendation', 'N/A')}\n\n"
    else:
        md += "*No significant conflicts detected*\n"

    md += "---\n\n## Security Considerations\n\n"
    for s in architecture.get('security_considerations', []):
        if isinstance(s, dict):
            md += f"- **{s.get('area', 'N/A')}** ({s.get('risk_level', 'medium')}): {s.get('description', 'N/A')}\n"
            md += f"  *Mitigation: {s.get('mitigation', 'N/A')}*\n\n"

    md += "---\n\n## Scalability Recommendations\n\n"
    for rec in architecture.get('scalability_recommendations', []):
        md += f"- {rec}\n"

    cost = architecture.get('cost_estimation', {})
    if isinstance(cost, dict) and cost:
        md += "\n---\n\n## Cost Estimation\n\n"
        monthly = cost.get('monthly_estimate', {})
        if monthly:
            md += "| Scale | Estimated Monthly Cost |\n|-------|----------------------|\n"
            for scale, est in monthly.items():
                md += f"| {scale.replace('_', ' ').title()} | {est} |\n"
        drivers = cost.get('major_cost_drivers', [])
        if drivers:
            md += "\n**Major Cost Drivers:**\n" + "".join(f"- {d}\n" for d in drivers)
        tips = cost.get('optimization_tips', [])
        if tips:
            md += "\n**Optimization Tips:**\n" + "".join(f"- {t}\n" for t in tips)

    md += "\n---\n\n## Architecture Diagrams\n\n"
    diagrams = architecture.get('mermaid_diagrams', {})
    diagram_titles = {
        'system_architecture': 'System Architecture', 'component_diagram': 'Component Diagram',
        'data_flow': 'Data Flow Diagram', 'deployment': 'Deployment Architecture', 'sequence': 'Sequence Diagram',
    }
    for key, title in diagram_titles.items():
        if key in diagrams:
            md += f"### {title}\n\n```mermaid\n{diagrams[key]}\n```\n\n"

    md += "\n---\n\n*Report generated by AI Architecture Designer*\n"
    return md


def generate_json_export(architecture: dict) -> str:
    """Generate a formatted JSON export of the architecture."""
    export_data = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "tool": "AI Architecture Designer",
            "version": "1.0.0",
        },
        "architecture": architecture,
    }
    return json.dumps(export_data, indent=2, ensure_ascii=False)


def generate_pdf_report(architecture: dict) -> bytes:
    """Generate a PDF report from architecture data using ReportLab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         Table, TableStyle, HRFlowable)
        from reportlab.lib.enums import TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=0.75 * inch,
                                 leftMargin=0.75 * inch, topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=22,
                                      textColor=colors.HexColor('#1e3a8a'), spaceAfter=10)
        h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=15,
                             textColor=colors.HexColor('#1e40af'), spaceBefore=14, spaceAfter=6)
        h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12,
                             textColor=colors.HexColor('#2563eb'), spaceBefore=10, spaceAfter=4)
        body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10,
                               textColor=colors.HexColor('#1f2937'), spaceAfter=5, leading=13)

        story = [
            Paragraph("Software Architecture Design Report", title_style),
            Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']),
            HRFlowable(width="100%", thickness=2, color=colors.HexColor('#2563eb')),
            Spacer(1, 10),
            Paragraph("Project Summary", h1),
            Paragraph(architecture.get('project_summary', 'N/A'), body),
            Spacer(1, 8),
            Paragraph("Recommended Architecture", h1),
            Paragraph(architecture.get('recommended_architecture', 'N/A'), body),
            Spacer(1, 8),
            Paragraph("Functional Requirements", h1),
        ]
        for req in architecture.get('functional_requirements', [])[:20]:
            story.append(Paragraph(f"&bull; {req}", body))
        story.append(Spacer(1, 8))

        tech_stack = architecture.get('tech_stack_recommendations', [])
        if tech_stack:
            story.append(Paragraph("Technology Stack", h1))
            table_data = [['Layer', 'Technology', 'Version', 'Justification']]
            for t in tech_stack[:15]:
                if isinstance(t, dict):
                    just = t.get('justification', 'N/A')
                    just = just[:60] + '...' if len(just) > 60 else just
                    table_data.append([t.get('layer', 'N/A'), t.get('technology', 'N/A'),
                                        t.get('version', 'N/A'), just])
            if len(table_data) > 1:
                table = Table(table_data, colWidths=[1.2 * inch, 1.5 * inch, 0.8 * inch, 3 * inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a8a')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#eff6ff')]),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#93c5fd')),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))

        conflicts = architecture.get('design_conflicts', [])
        if conflicts:
            story.append(Paragraph("Design Conflicts & Risks", h1))
            for c in conflicts[:10]:
                if isinstance(c, dict):
                    story.append(Paragraph(
                        f"[{c.get('severity', 'medium').upper()}] {c.get('type', 'Issue').replace('_', ' ').title()}", h2))
                    story.append(Paragraph(c.get('description', 'N/A'), body))
                    story.append(Paragraph(f"<b>Recommendation:</b> {c.get('recommendation', 'N/A')}", body))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#93c5fd')))
        story.append(Paragraph("Generated by AI Architecture Designer",
                                ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8,
                                               textColor=colors.grey, alignment=TA_CENTER)))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        return generate_markdown_report(architecture).encode('utf-8')


def generate_docx_report(architecture: dict) -> bytes:
    """Generate a DOCX report from architecture data using python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading('Software Architecture Design Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_heading('Project Summary', level=1)
        doc.add_paragraph(architecture.get('project_summary', 'N/A'))

        doc.add_heading('Recommended Architecture', level=1)
        doc.add_paragraph(architecture.get('recommended_architecture', 'N/A'))

        doc.add_heading('Functional Requirements', level=1)
        for req in architecture.get('functional_requirements', []):
            doc.add_paragraph(req, style='List Bullet')

        doc.add_heading('Non-Functional Requirements', level=1)
        for req in architecture.get('non_functional_requirements', []):
            doc.add_paragraph(req, style='List Bullet')

        doc.add_heading('System Components', level=1)
        for comp in architecture.get('components', []):
            if isinstance(comp, dict):
                doc.add_heading(comp.get('name', 'Unknown'), level=2)
                p = doc.add_paragraph(); p.add_run('Type: ').bold = True; p.add_run(comp.get('type', 'N/A'))
                p = doc.add_paragraph(); p.add_run('Technology: ').bold = True; p.add_run(comp.get('technology', 'N/A'))
                p = doc.add_paragraph(); p.add_run('Description: ').bold = True; p.add_run(comp.get('description', 'N/A'))

        doc.add_heading('Technology Stack', level=1)
        tech_stack = architecture.get('tech_stack_recommendations', [])
        if tech_stack:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, h in enumerate(['Layer', 'Technology', 'Version', 'Justification']):
                hdr[i].text = h
            for t in tech_stack:
                if isinstance(t, dict):
                    row = table.add_row()
                    row.cells[0].text = t.get('layer', 'N/A')
                    row.cells[1].text = t.get('technology', 'N/A')
                    row.cells[2].text = t.get('version', 'N/A')
                    row.cells[3].text = t.get('justification', 'N/A')

        doc.add_heading('Design Conflicts & Risks', level=1)
        for c in architecture.get('design_conflicts', []):
            if isinstance(c, dict):
                doc.add_heading(f"{c.get('severity', 'N/A').upper()}: {c.get('type', 'Issue')}", level=2)
                doc.add_paragraph(c.get('description', 'N/A'))
                p = doc.add_paragraph(); p.add_run('Recommendation: ').bold = True
                p.add_run(c.get('recommendation', 'N/A'))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        return generate_markdown_report(architecture).encode('utf-8')


def format_component_count(architecture: dict) -> dict:
    """Extract counts used for dashboard metric cards."""
    components = architecture.get('components', [])
    conflicts = architecture.get('design_conflicts', [])
    tech_stack = architecture.get('tech_stack_recommendations', [])
    security = architecture.get('security_considerations', [])

    high_conflicts = [c for c in conflicts if isinstance(c, dict) and c.get('severity') == 'high']
    high_security = [s for s in security if isinstance(s, dict) and s.get('risk_level') == 'high']

    return {
        'total_components': len(components),
        'total_conflicts': len(conflicts),
        'high_severity_conflicts': len(high_conflicts),
        'tech_stack_count': len(tech_stack),
        'security_issues': len(security),
        'high_security_risks': len(high_security),
        'functional_reqs': len(architecture.get('functional_requirements', [])),
        'non_functional_reqs': len(architecture.get('non_functional_requirements', [])),
    }
